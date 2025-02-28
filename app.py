import streamlit as st
import whisper
import torch
import tempfile
import os
from datetime import datetime
import google.generativeai as genai
from docx import Document
from keybert import KeyBERT

# Retrieve API key from environment variable
api_key = os.getenv("GEMINI_API_KEY")
if not api_key:
    st.error("Gemini API key not found. Please set it in your environment variables.")
    st.stop()

# Configure Gemini API
genai.configure(api_key=api_key)

# Load Whisper model for speech-to-text
@st.cache_resource
def load_whisper_model():
    device = "cuda" if torch.cuda.is_available() else "cpu"
    return whisper.load_model("base").to(device)

# Transcribe audio
def transcribe_audio(audio_path):
    try:
        model = load_whisper_model()
        result = model.transcribe(audio_path)
        return result["text"]
    except Exception as e:
        st.error(f"Error in transcription: {e}")
        return None

# Perform speaker diarization using Gemini (since we removed Hugging Face)
def assign_speakers_with_gemini(transcribed_text):
    try:
        prompt = f"""
        The following is a transcription of a meeting. Identify different speakers and format it clearly.
        Assign labels like 'Speaker 1:', 'Speaker 2:', etc.

        --- Transcription Start ---
        {transcribed_text}
        --- Transcription End ---

        Return a well-structured conversation format with correct speaker assignments.
        """
        response = genai.GenerativeModel("gemini-pro").generate_content(prompt)
        return response.text.strip()
    except Exception as e:
        st.error(f"Error in speaker diarization: {e}")
        return "Error: Speaker identification failed."

# Summarize text using Gemini
def summarize_text(text, summary_length=5):
    try:
        prompt = f"Summarize the following meeting transcript in {summary_length} sentences:\n{text}"
        response = genai.GenerativeModel("gemini-pro").generate_content(prompt)
        return response.text.strip()
    except Exception as e:
        st.error(f"Error summarizing text: {e}")
        return "Error: Summarization failed."

# Analyze sentiment
def analyze_sentiment(text):
    try:
        prompt = f"Analyze the sentiment of the following meeting:\n{text}\nClassify as 'Positive', 'Negative', or 'Neutral'."
        response = genai.GenerativeModel("gemini-pro").generate_content(prompt)
        return response.text.strip()
    except Exception as e:
        st.error(f"Error analyzing sentiment: {e}")
        return "Error: Sentiment analysis failed."

# Extract keywords
def extract_keywords(text):
    try:
        kw_model = KeyBERT()
        keywords = kw_model.extract_keywords(text, keyphrase_ngram_range=(1, 2), stop_words='english', top_n=5)
        return [word[0] for word in keywords]
    except Exception as e:
        st.error(f"Error extracting keywords: {e}")
        return []

# Detect action items
def extract_action_items(text):
    try:
        prompt = f"Extract the action items from the following meeting transcript:\n{text}"
        response = genai.GenerativeModel("gemini-pro").generate_content(prompt)
        return response.text.strip()
    except Exception as e:
        st.error(f"Error extracting action items: {e}")
        return "Error: Could not extract action items."

# Save minutes as a Word document
def save_as_docx(diarized_text, summary, sentiment, keywords, action_items, filename="minutes_of_meeting.docx"):
    try:
        doc = Document()
        doc.add_heading("Minutes of Meeting", 0)

        doc.add_heading("Speaker Diarization", level=1)
        doc.add_paragraph(diarized_text)

        doc.add_heading("Summary", level=1)
        doc.add_paragraph(summary)

        doc.add_heading("Sentiment Analysis", level=1)
        doc.add_paragraph(f"Sentiment: {sentiment}")

        doc.add_heading("Keywords", level=1)
        doc.add_paragraph(", ".join(keywords))

        doc.add_heading("Action Items", level=1)
        doc.add_paragraph(action_items)

        doc.save(filename)
        return filename
    except Exception as e:
        st.error(f"Error saving Word document: {e}")
        return None

# Streamlit App
def main():
    st.title("📜 Meeting Minutes Generator")
    st.write("Upload a meeting recording to generate structured minutes.")

    uploaded_file = st.file_uploader("Upload Audio File", type=["mp3", "wav", "m4a"])

    if uploaded_file:
        st.audio(uploaded_file, format="audio/wav")

        try:
            with tempfile.NamedTemporaryFile(delete=False, suffix=".wav") as tmp_file:
                tmp_file.write(uploaded_file.read())
                audio_path = tmp_file.name

            # Transcription
            st.write("Transcribing audio...")
            with st.spinner("Processing transcription..."):
                transcribed_text = transcribe_audio(audio_path)
            if not transcribed_text:
                return

            st.subheader("Transcription")
            st.write(transcribed_text)

            # Speaker Diarization
            st.write("Identifying speakers...")
            with st.spinner("Processing speaker diarization..."):
                diarized_text = assign_speakers_with_gemini(transcribed_text)

            st.subheader("Speaker Diarization")
            st.write(diarized_text)

            # Summary
            summary_length = st.slider("Summary Length (sentences)", 1, 10, 10)
            with st.spinner("Generating summary..."):
                summary = summarize_text(diarized_text, summary_length)

            st.subheader("Summary")
            st.write(summary)

            # Sentiment Analysis
            with st.spinner("Analyzing sentiment..."):
                sentiment = analyze_sentiment(transcribed_text)

            st.subheader("Sentiment Analysis")
            st.write(f"Sentiment: {sentiment}")

            # Keyword Extraction
            with st.spinner("Extracting keywords..."):
                keywords = extract_keywords(transcribed_text)

            st.subheader("Keywords")
            st.write(", ".join(keywords))

            # Action Items
            with st.spinner("Extracting action items..."):
                action_items = extract_action_items(transcribed_text)

            st.subheader("Action Items")
            st.write(action_items)

            # Save as Word Document
            with st.spinner("Saving as Word document..."):
                docx_filename = save_as_docx(diarized_text, summary, sentiment, keywords, action_items)
                if docx_filename:
                    with open(docx_filename, "rb") as file:
                        st.download_button("📥 Download Minutes", data=file, file_name=docx_filename, mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

        except Exception as e:
            st.error(f"An error occurred: {e}")

        finally:
            if os.path.exists(audio_path):
                os.remove(audio_path)

if __name__ == "__main__":
    main()
